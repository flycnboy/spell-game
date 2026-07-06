from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date

doc = Document()

# 标题
title = doc.add_heading('梓晏的单词学习 — 操作与运维手册', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(f'版本日期：{date.today().isoformat()}\n适用人群：家长、运营维护人员', style='Subtitle')

# ===== 一、系统概述 =====
doc.add_heading('一、系统概述', 1)
doc.add_paragraph(
    '"梓晏的单词学习"是一款面向小学生的英语单词学习应用（PWA），支持安卓手机和平板。\n'
    '• 两种学习模式：拼写模式（看打散字母拼单词）和听写模式（听发音直接输入）。\n'
    '• 词库管理：支持远程同步（Gist）、手动创建、批量导入/导出。\n'
    '• 智能复习：基于艾宾浩斯遗忘曲线自动安排每日学习计划。\n'
    '• 单词释义：联网获取中英文释义和例句。\n'
    '• 统计功能：记录正确率、错题本及强化练习。\n'
    '• 无需安装：浏览器打开即可使用，可添加到手机主屏幕。'
)

# ===== 二、快速开始 =====
doc.add_heading('二、快速开始', 1)

doc.add_heading('2.1 在手机上使用（推荐）', 2)
doc.add_paragraph(
    '1. 打开手机 Chrome 浏览器。\n'
    '2. 访问线上版：https://flycnboy.github.io/spell-game/\n'
    '3. 点击 Chrome 菜单（右上角三个点）→ "添加到主屏幕"。\n'
    '4. 桌面出现图标，点击即可像 App 一样使用。'
)

doc.add_heading('2.2 在电脑上开发/测试', 2)
doc.add_paragraph(
    '1. 打开终端，进入项目目录：cd D:/SynologyDrive/SynologyDrive/MyProject/spell-game\n'
    '2. 启动开发服务器：npm run dev\n'
    '3. 浏览器打开 http://localhost:5173/spell-game/'
)

doc.add_heading('2.3 前置条件 — 语音引擎', 2)
doc.add_paragraph(
    '• 安卓设备需要安装 Google 文字转语音引擎（Google TTS）。\n'
    '• 如果喇叭没有声音，请检查：设置 → 语言和输入法 → 文字转语音(TTS) → 选择 Google TTS 引擎。\n'
    '• 注意：仅在 HTTPS（线上版）下 Web Speech API 才能工作，HTTP 本地开发不会发音。'
)

# ===== 三、板块详解 =====
doc.add_heading('三、功能板块详解', 1)

doc.add_heading('3.1 学习模式', 2)
doc.add_paragraph(
    '• 拼写模式：孩子看到打散的字母块，点击字母拼出单词，按确认提交。\n'
    '• 听写模式：只有发音，没有字母提示。孩子直接在输入框里写出听到的单词。\n'
    '• 流程：听音 → 拼写/听写 → 判断对错 → 下一题 → 完成总结。\n'
    '• 键盘快捷键（电脑端）：字母键选字、Enter 确认、Backspace 删除最后一个字母。'
)

doc.add_heading('3.2 单词录入与管理', 2)
p = doc.add_paragraph()
p.add_run('支持四种录入方式：\n').bold = True
doc.add_paragraph(
    '① 同步远程：点 "🔄 同步远程"，粘贴 Gist Raw URL（默认已填好最新词库）。\n'
    '② 新建本地：点 "＋ 新建本地"，填写名称和单词列表（一行一个）。\n'
    '③ 导入 JSON：点 "📥 导入文件"，选择之前导出的 .json 文件合并。\n'
    '④ 导出全部：点 "📤 导出全部"，下载所有词库为一个 .json 文件'
)

doc.add_heading('3.3 词库详情编辑', 2)
doc.add_paragraph(
    '点击每个词库右侧的 "📝" 进入编辑页，可以：\n'
    '• 增删单词、批量输入（逗号或换行分隔）。\n'
    '• 从其他词库导入单词（勾选即可）。\n'
    '• 勾选特定单词 → 另存为生词库（如 "超难词汇"）。\n'
    '• 点击单词标签可删除，或通过输入框新增。\n'
    '• 导出该词库为含释义的完整 JSON。\n\n'
    '词库名称修改：点击 "✏️" 直接内联改名。'
)

doc.add_heading('3.4 单词释义（📖）', 2)
doc.add_paragraph(
    '• 选中一个词库后点 "📖 释义"，自动为每个单词获取中文翻译和英文例句。\n'
    '• 调用免费 API（Dictionary API + MyMemory），无请求次数限制。\n'
    '• 完成后缓存到本地，下次不再消耗流量。\n'
    '• 释义会在听音、拼写、结果页面展示。'
)

doc.add_heading('3.5 学习计划（艾宾浩斯遗忘曲线）', 2)
doc.add_paragraph(
    '• 设置："⚙️ 设置" → 每天学几个新词、计划多少天。\n'
    '• 新单词在第 1、7、16、30 天后自动安排复习。\n'
    '• 每天练习时，先展示 "今日计划"，列出新词 + 历史复习词。'
)

doc.add_heading('3.6 统计与成果', 2)
doc.add_paragraph(
    '• 总学习次数、正确率、今日学习量。\n'
    '• 错词列表（支持强化练习，答对自动移除）。\n'
    '• "强化练习错词"功能：只出错的词，直到全部掌握为止。\n'
    '• 清零重置：清空所有历史统计。'
)

doc.add_heading('3.7 完成页面', 2)
doc.add_paragraph(
    '• 显示本轮正确/错误数量、正确率。\n'
    '• 三颗星评级（正确率 ≥ 90% 三颗）。\n'
    '• "再来一轮" 或 "强化练习错词"。\n'
    '• 提醒：每天练习结束自动推进到第二天，复习会按曲线自动生成。'
)

# ===== 四、运维与部署 =====
doc.add_heading('四、运维与部署', 1)

doc.add_heading('4.1 更新单词库（Gist 同步）', 2)
doc.add_paragraph(
    '1. 电脑上打开你的 Gist：https://gist.github.com/flycnboy/86d67f60c66736f6dd6092edceab5edf\n'
    '2. 点击 "Edit" 编辑内容，增删单词。\n'
    '3. 点击 "Update gist" 保存。\n'
    '4. 手机上打开 App → 🔄 同步远程 → 同步。'
)

doc.add_heading('4.2 代码部署流程（从开发到发布）', 2)
doc.add_paragraph(
    '1. 本地修改代码并验证：npm run dev（浏览器打开 http://localhost:5173/spell-game/）\n'
    '2. 构建生产版本：npm run build（输出到 docs/）\n'
    '3. 提交并推送到 Gitee：\n'
    '   git add .\n'
    '   git commit -m "描述本次改动"\n'
    '   git push origin main\n'
    '4. 在 Gitee 管理 → 仓库镜像 → 手动同步到 GitHub\n'
    '5. 等待 1-2 分钟，GitHub Pages 自动更新。'
)

doc.add_heading('4.3 项目结构', 2)
p = doc.add_paragraph()
p.add_run('核心源代码目录（src/）：\n').bold = True
doc.add_paragraph(
    '• src/App.tsx — 主应用入口，路由控制\n'
    '• src/components/ — 所有页面组件\n'
    '• src/hooks/ — 自定义 hook：语音、统计、词库、艾宾浩斯、释义\n'
    '• src/store/gameStore.ts — Zustand 状态管理\n'
    '• src/types.ts — 类型定义\n'
    '• 部署产物目录：docs/（推送到 GitHub Pages 的静态文件）'
)

doc.add_heading('4.4 线上地址', 2)
doc.add_paragraph(
    '• 正式版：https://flycnboy.github.io/spell-game/\n'
    '• Gitee 仓库：https://gitee.com/flycnboy/spell-game\n'
    '• GitHub 仓库：https://github.com/flycnboy/spell-game'
)

# ===== 五、常见问题 =====
doc.add_heading('五、常见问题', 1)

faqs = [
    ('喇叭没有声音怎么办？',
     '① 确保已安装 Google 文字转语音引擎（TTS）。\n'
     '② 确保使用线上 HTTPS 版（不是本地 HTTP 开发版）。\n'
     '③ 确保手机媒体音量已打开。'
    ),
    ('释义显示为空？',
     '需要先在词库管理页点 "📖 释义"，为当前词库联网获取一次数据。获取后数据缓存在本地，以后不需要再联网。'
    ),
    ('同步失败怎么办？',
     '检查 Gist URL 是否正确（必须以 raw.githubusercontent.com 或 gist.githubusercontent.com 开头）、是否 Public（公开）、是否确实有内容。'
    ),
    ('数据丢失怎么办？',
     '定期使用 "📤 导出全部" 备份词库。如果确实丢失，可通过 "📥 导入文件" 恢复。'
    ),
    ('如何给孩子的平板添加 App？',
     '打开线上版 → Chrome 菜单 → "添加到主屏幕"。桌面会出现独立图标，点击后全屏打开。'
    ),
]
for q, a in faqs:
    p = doc.add_paragraph()
    p.add_run(f'Q: {q}\n').bold = True
    p.add_run(f'A: {a}')

doc.save('梓晏的单词学习_操作手册.docx')
print('文档已保存：梓晏的单词学习_操作手册.docx')
